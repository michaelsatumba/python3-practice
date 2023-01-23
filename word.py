import docx
import os

os.chdir('/Users/michaelsatumba/Desktop/')
print(os.getcwd())

'''
d = docx.Document('/Users/michaelsatumba/Desktop/master resume.docx')
#print(d.paragraphs)
#print(d.paragraphs[0].text)
p = d.paragraphs[0]
#print(p.text)
print(p.text)
p.runs[2].underline = True
p.runs[2].text = 'AHHH'
print(p.style)
p.style = 'normal'
print(p.style)
d.save('master resume.docx')
'''

'''
d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
d.save('demo4.docx')
'''

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
print(getText('/Users/michaelsatumba/Desktop/master resume.docx'))
